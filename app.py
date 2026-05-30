# -*- coding: utf-8 -*-
"""
Générateur de cartels à partir d'un fichier Excel musée
- lecture des colonnes sur la 2e ligne (header=1)
- choix de l'onglet : Liste, Repro, ou les deux
- aperçu des données
- choix des champs à afficher via cases à cocher
- mise en page fixe des cartels
- gestion robuste des doublons de colonnes
- formatage du "e" en exposant dans le champ Date (ex: XIXe siècle)
- uniformisation des apostrophes
- export Word
"""

import io
import re
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# --------------------------------------------------
# Configuration de la page
# --------------------------------------------------
st.set_page_config(
    page_title="Générateur de cartels",
    page_icon="logo_musee.png",
    layout="centered"
)

col1, col2 = st.columns([1, 6])

with col1:
    st.image("logo_musee.png", width=85)

with col2:
    st.markdown(
        "<h1 style='margin-top: 10px; margin-bottom: 0;'>Générateur de cartels</h1>",
        unsafe_allow_html=True
    )
    
st.write(
    "Cette interface permet de générer des cartels à partir d’un fichier Excel fourni par le musée."
)

st.markdown(
    "<h4 style='text-align: center;'> Insérer votre fichier Excel</h4>",
    unsafe_allow_html=True
)

uploaded = st.file_uploader("", type=["xlsx", "xls"])


# --------------------------------------------------
# Constantes
# --------------------------------------------------
ALL_FIELDS = [
    "Auteur / Exécutant",
    "Date auteur/exécutant",
    "Editeur",
    "Titre",
    "Provenance",
    "Date",
    "Technique(s) de l'œuvre originale",
    "Mention Collection pour le cartel",
    "Information sur l'acquisition",
]

REQUIRED_FIELDS = [
    "Titre",
    "Technique(s) de l'œuvre originale",
    "Mention Collection pour le cartel",
]

EXPECTED_SHEETS = ["Liste", "Repro"]


# --------------------------------------------------
# Fonctions utilitaires
# --------------------------------------------------
def add_horizontal_rule(doc):
    """Ajoute une ligne horizontale comme séparateur."""
    p = doc.add_paragraph()
    p_par = p._p
    pPr = p_par.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def clean_filename(name):
    """Nettoie le nom du fichier."""
    name = str(name).strip()
    name = re.sub(r'[\\/*?:"<>|]', "_", name)
    return name


def normalize_apostrophes(text):
    """
    Remplace les variantes d'apostrophes par une apostrophe simple.
    """
    if not isinstance(text, str):
        return text

    return (
        text.replace("’", "'")
            .replace("ʼ", "'")
            .replace("`", "'")
            .replace("´", "'")
    )


def make_unique_columns(columns):
    """
    Rend les noms de colonnes uniques.
    Exemple :
    ['Titre', 'Titre'] -> ['Titre', 'Titre__2']
    """
    seen = {}
    new_cols = []

    for col in columns:
        col = str(col).strip()
        if col in seen:
            seen[col] += 1
            new_cols.append(f"{col}__{seen[col]}")
        else:
            seen[col] = 1
            new_cols.append(col)

    return new_cols


def normalize_columns(df):
    """Nettoie et rend uniques les noms de colonnes."""
    df.columns = make_unique_columns(df.columns)
    return df


def safe(val):
    """
    Renvoie une chaîne propre, même si val est une Series ou une liste.
    Uniformise aussi les apostrophes.
    """
    if val is None:
        return ""

    if isinstance(val, pd.Series):
        vals = [safe(v) for v in val.tolist()]
        vals = [v for v in vals if v]
        return " | ".join(vals)

    if isinstance(val, (list, tuple)):
        vals = [safe(v) for v in val]
        vals = [v for v in vals if v]
        return " | ".join(vals)

    try:
        if pd.isna(val):
            return ""
    except Exception:
        pass

    # Si c'est un nombre entier lu comme décimal, ex : 5.0 -> 5
    if isinstance(val, float) and val.is_integer():
        return str(int(val))
        
    return normalize_apostrophes(str(val).strip())


def get_cell_value(row, column_name):
    """Récupère une valeur de cellule de façon robuste."""
    if column_name in row.index:
        return safe(row[column_name])
    return ""


def read_sheet(uploaded_file, sheet_name):
    """
    Lit un onglet Excel en considérant que :
    - la 1ère ligne n'est pas à prendre en compte
    - les en-têtes sont sur la 2ème ligne
    """
    df = pd.read_excel(uploaded_file, sheet_name=sheet_name, header=1)
    df = normalize_columns(df)
    return df


def validate_required_columns(df, required_fields):
    """Vérifie la présence des colonnes obligatoires exactes."""
    missing = [col for col in required_fields if col not in df.columns]
    return missing


def drop_empty_rows(df, required_fields):
    """Supprime les lignes sans contenu utile sur les champs obligatoires."""
    working_df = df.copy()

    for col in required_fields:
        if col in working_df.columns:
            working_df[col] = working_df[col].apply(safe)

    mask = working_df[required_fields].apply(
        lambda row: any(val != "" for val in row),
        axis=1
    )
    return working_df[mask].reset_index(drop=True)


def add_date_with_superscript(paragraph, text, font_size=Pt(10.5)):
    """
    Met en exposant le 'e' uniquement pour les siècles en chiffres romains :
    XIXe siècle, XXe siècle, fin du XIXe siècle, etc.
    Sans toucher au 'e' de 'siècle'.
    """
    pattern = re.compile(r'(?i)\b([IVXLCDM]+)(e)(?=\s|[-–—,;:.!?)]|$)')
    last_idx = 0

    for match in pattern.finditer(text):
        start, end = match.span()

        if start > last_idx:
            run = paragraph.add_run(text[last_idx:start])
            run.font.size = font_size

        run_roman = paragraph.add_run(match.group(1))
        run_roman.font.size = font_size

        run_e = paragraph.add_run(match.group(2))
        run_e.font.size = font_size
        run_e.font.superscript = True

        last_idx = end

    if last_idx < len(text):
        run = paragraph.add_run(text[last_idx:])
        run.font.size = font_size


# --------------------------------------------------
# Cartels
# --------------------------------------------------
def add_cartel_to_doc(doc, row, selected_fields, source_sheet=None):
    """
    Ordre d'affichage :
    1. Titre
    2. Auteur / Exécutant, Date auteur/exécutant
    3. Editeur
    4. Provenance
    5. Date
    6. Technique(s) de l'œuvre originale
    7. Mention Collection pour le cartel
    8. Information sur l’acquisition
    """
    numero_expo = get_cell_value(row, "n°expo")
    type_cartel = get_cell_value(row, "Type de cartel")
    
    titre = get_cell_value(row, "Titre") or "Sans titre"

    auteur = get_cell_value(row, "Auteur / Exécutant")
    date_auteur = get_cell_value(row, "Date auteur/exécutant")
    editeur = get_cell_value(row, "Editeur")
    provenance = get_cell_value(row, "Provenance")
    date_oeuvre = get_cell_value(row, "Date")
    technique = get_cell_value(row, "Technique(s) de l'œuvre originale")
    mention = get_cell_value(row, "Mention Collection pour le cartel")
    acquisition = get_cell_value(row, "Information sur l'acquisition")

    # n°expo affiché automatiquement au-dessus du cartel
    if numero_expo:
        p_expo = doc.add_paragraph()
        p_expo.paragraph_format.space_after = Pt(1)
        r_expo = p_expo.add_run(numero_expo)
        r_expo.italic = True
        r_expo.font.size = Pt(9)
        r_expo.font.color.rgb = RGBColor(160, 160, 160)

        if type_cartel.strip().lower() == "développé":
            r_expo.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # 1. Titre
    p_titre = doc.add_paragraph()
    p_titre.paragraph_format.space_after = Pt(2)
    r_titre = p_titre.add_run(titre)
    r_titre.bold = True
    r_titre.font.size = Pt(14)

    # 2. Auteur / Exécutant, Date auteur/exécutant
    auteur_parts = []
    if "Auteur / Exécutant" in selected_fields and auteur:
        auteur_parts.append(auteur)
    if "Date auteur/exécutant" in selected_fields and date_auteur:
        auteur_parts.append(date_auteur)

    if auteur_parts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(", ".join(auteur_parts))
        r.font.size = Pt(10.5)

    # 3. Editeur
    if "Editeur" in selected_fields and editeur:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(editeur)
        r.font.size = Pt(10.5)

    # 4. Provenance
    if "Provenance" in selected_fields and provenance:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(provenance)
        r.font.size = Pt(10.5)

    # 5. Date
    if "Date" in selected_fields and date_oeuvre:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_date_with_superscript(p, date_oeuvre, font_size=Pt(10.5))

    # 6. Technique(s) de l'œuvre originale
    if "Technique(s) de l'œuvre originale" in selected_fields and technique:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(technique)
        r.font.size = Pt(10.5)

    # 7. Mention Collection pour le cartel
    if "Mention Collection pour le cartel" in selected_fields and mention:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(mention)
        r.font.size = Pt(10.5)

    # 8. Information sur l’acquisition
    if "Information sur l'acquisition" in selected_fields and acquisition:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(acquisition)
        r.font.size = Pt(10.5)

    # Affichage optionnel de l'onglet source si un seul onglet sélectionné
    if source_sheet:
        p_sheet = doc.add_paragraph()
        p_sheet.paragraph_format.space_before = Pt(3)
        r_sheet = p_sheet.add_run(f"Onglet : {source_sheet}")
        r_sheet.italic = True
        r_sheet.font.size = Pt(8.5)


def create_word_document(data_by_sheet, selected_fields, document_title):
    """Crée le document Word final."""
    doc = Document()
    doc.core_properties.title = document_title

    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(document_title)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    sheet_names = list(data_by_sheet.keys())

    for sheet_index, sheet_name in enumerate(sheet_names):
        df = data_by_sheet[sheet_name]

        if len(sheet_names) > 1:
            p_sheet_title = doc.add_paragraph()
            r_sheet_title = p_sheet_title.add_run(sheet_name)
            r_sheet_title.bold = True
            r_sheet_title.font.size = Pt(16)
            doc.add_paragraph()

        for row_index, (_, row) in enumerate(df.iterrows()):
            add_cartel_to_doc(
                doc=doc,
                row=row,
                selected_fields=selected_fields,
                source_sheet=sheet_name if len(sheet_names) == 1 else None
            )

            is_last_row_of_sheet = row_index == len(df) - 1
            is_last_sheet = sheet_index == len(sheet_names) - 1

            if not (is_last_row_of_sheet and is_last_sheet):
                doc.add_paragraph()
                add_horizontal_rule(doc)
                doc.add_paragraph()

        if len(sheet_names) > 1 and not is_last_sheet:
            doc.add_paragraph()
            doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --------------------------------------------------
# Interface principale
# --------------------------------------------------
if uploaded:
    try:
        excel_file = pd.ExcelFile(uploaded)
        available_sheets = excel_file.sheet_names
    except Exception as e:
        st.error(f"Erreur de lecture du fichier Excel : {e}")
        st.stop()

    st.subheader("Paramètres")

    selectable_sheets = [s for s in EXPECTED_SHEETS if s in available_sheets]

    if not selectable_sheets:
        st.error("Les onglets attendus ('Liste' et/ou 'Repro') n'ont pas été trouvés dans le fichier.")
        st.stop()

    selected_sheets = st.multiselect(
        "Choisir le ou les onglets à utiliser",
        options=selectable_sheets,
        default=selectable_sheets
    )

    if not selected_sheets:
        st.warning("Veuillez sélectionner au moins un onglet.")
        st.stop()

    # Aperçu sur le premier onglet sélectionné
    preview_sheet = selected_sheets[0]
    try:
        preview_df = read_sheet(uploaded, preview_sheet)
    except Exception as e:
        st.error(f"Erreur de lecture de l'onglet '{preview_sheet}' : {e}")
        st.stop()

    if preview_df.empty:
        st.warning(f"L'onglet '{preview_sheet}' est vide.")
    else:
        st.subheader(f"Aperçu des données — onglet {preview_sheet}")
        st.dataframe(preview_df.head(10), use_container_width=True)

    invalid_sheets = {}
    loaded_sheets = {}

    for sheet in selected_sheets:
        try:
            df_sheet = read_sheet(uploaded, sheet)
            loaded_sheets[sheet] = df_sheet
            missing = validate_required_columns(df_sheet, REQUIRED_FIELDS)
            if missing:
                invalid_sheets[sheet] = missing
        except Exception as e:
            st.error(f"Erreur de lecture de l'onglet '{sheet}' : {e}")
            st.stop()

    if invalid_sheets:
        for sheet, missing in invalid_sheets.items():
            st.error(
                f"Onglet '{sheet}' : colonnes obligatoires manquantes : {', '.join(missing)}"
            )
        st.stop()

    available_fields = []
    for field in ALL_FIELDS:
        if any(field in df.columns for df in loaded_sheets.values()):
            available_fields.append(field)

    st.markdown("### Champs à afficher sur les cartels")

    selected_fields = st.multiselect(
        "Sélectionner les champs à faire apparaître",
        options=available_fields,
        default=[field for field in REQUIRED_FIELDS if field in available_fields]
    )

    missing_required_in_selection = [f for f in REQUIRED_FIELDS if f not in selected_fields]
    if missing_required_in_selection:
        st.warning(
            "Les champs obligatoires seront toujours inclus : "
            + ", ".join(missing_required_in_selection)
        )
        selected_fields = list(dict.fromkeys(selected_fields + missing_required_in_selection))

    st.markdown(
        "<div style='font-weight:600; margin-bottom:4px;'>"
        "Nom du fichier Word (sans extension) <span style='color:#d00'>*</span>"
        "</div>",
        unsafe_allow_html=True
    )

    nom_fichier = st.text_input(
        label="",
        placeholder="Indiquer le nom du document",
        key="nom_fichier_input",
        label_visibility="collapsed"
    )

    titre_document = "Cartels"

    if st.button("Transformer"):
        if not nom_fichier.strip():
            st.error("Veuillez indiquer le nom du document.")
            st.markdown("""
                <style>
                div[data-testid="stTextInput"] input {
                    border: 1px solid #d00 !important;
                    box-shadow: 0 0 0 1px rgba(221,0,0,.25) !important;
                }
                </style>
            """, unsafe_allow_html=True)
            st.stop()

        nom_fichier = clean_filename(nom_fichier)
        if not nom_fichier:
            st.error("Le nom du fichier n'est pas valide.")
            st.stop()

        export_data = {}

        for sheet_name, df_sheet in loaded_sheets.items():
            filtered_df = drop_empty_rows(df_sheet, REQUIRED_FIELDS)

            if filtered_df.empty:
                st.warning(f"Aucune ligne exploitable trouvée dans l'onglet '{sheet_name}'.")
                continue

            cols_to_keep = [col for col in ALL_FIELDS if col in filtered_df.columns]
            
            extra_cols = []
            if "n°expo" in filtered_df.columns:
                extra_cols.append("n°expo")

            if "Type de cartel" in filtered_df.columns:
                extra_cols.append("Type de cartel")

            cols_to_keep = extra_cols + cols_to_keep
                
            export_data[sheet_name] = filtered_df[cols_to_keep].copy()

        if not export_data:
            st.error("Aucune donnée exploitable n'a été trouvée pour générer les cartels.")
            st.stop()

        try:
            buffer = create_word_document(
                data_by_sheet=export_data,
                selected_fields=selected_fields,
                document_title=titre_document
            )
        except Exception as e:
            st.error(f"Erreur lors de la génération du document Word : {e}")
            st.stop()

        total_cartels = sum(len(df) for df in export_data.values())

        st.success(f"Document généré avec succès ({total_cartels} cartel(s)) !")
        st.download_button(
            label="Télécharger le fichier Word",
            data=buffer,
            file_name=f"{nom_fichier}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
